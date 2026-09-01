# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "marimo>=0.13.0",
#     "pandas>=2.0.0",
#     "numpy>=1.24.0",
#     "matplotlib>=3.7.0",
#     "scikit-learn>=1.3.0",
#     "requests>=2.31.0",
#     "openpyxl>=3.1.0",
#     "python-docx>=1.1.0",
# ]
# ///

import marimo

__generated_with = "0.13.0"
app = marimo.App(width="full")


@app.cell
def _():
    import io
    from pathlib import Path
    from tempfile import TemporaryDirectory

    import marimo as mo
    import matplotlib.pyplot as plt
    import numpy as np
    import pandas as pd
    import requests

    from sklearn.decomposition import PCA
    from sklearn.linear_model import LinearRegression
    from sklearn.metrics import mean_absolute_percentage_error, r2_score
    from sklearn.preprocessing import StandardScaler

    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    pd.set_option("display.max_columns", 100)
    pd.set_option("display.width", 180)
    return (
        Document,
        Inches,
        LinearRegression,
        OxmlElement,
        PCA,
        Path,
        Pt,
        StandardScaler,
        TemporaryDirectory,
        WD_ALIGN_PARAGRAPH,
        WD_CELL_VERTICAL_ALIGNMENT,
        WD_TABLE_ALIGNMENT,
        io,
        mean_absolute_percentage_error,
        mo,
        np,
        pd,
        plt,
        qn,
        r2_score,
        requests,
    )


@app.cell
def _(mo):
    mo.md(
        r"""
        # ERAU Central Chiller Plant Utility Performance Analysis
        **North TES + South TES + Open-Meteo + PCA**

        This marimo app converts monthly/billing-period utility data into North and South plant totals, downloads hourly historical weather, builds weather-normalized baselines, flags unusual periods, performs PCA on weather regimes, and creates Excel and Word outputs.

        **Important:** This is a utility- and weather-based screening analysis. It does not calculate actual chiller kW/ton because measured plant kW, chilled-water flow, and cooling tonnage are not part of the current input.
        """
    )
    return


@app.cell
def _(mo):
    utility_upload = mo.ui.file(
        filetypes=[".csv", ".xlsx", ".xls"],
        multiple=False,
        label="Upload chiller utility import file",
    )
    latitude = mo.ui.number(value=29.2108, step=0.0001, label="Latitude")
    longitude = mo.ui.number(value=-81.0228, step=0.0001, label="Longitude")
    cdd_base = mo.ui.number(value=65.0, step=1.0, label="CDD base (F)")
    cooling_threshold = mo.ui.number(value=75.0, step=1.0, label="Cooling threshold (F)")
    humidity_threshold = mo.ui.number(value=70.0, step=1.0, label="Humidity threshold (%)")
    anomaly_z = mo.ui.number(value=1.5, step=0.1, label="Anomaly z threshold")
    blended_rate = mo.ui.number(value=0.12, step=0.001, label="Blended electric rate ($/kWh)")

    controls = mo.vstack(
        [
            mo.md("## 1. Input and assumptions"),
            utility_upload,
            mo.hstack([latitude, longitude, cdd_base]),
            mo.hstack([cooling_threshold, humidity_threshold, anomaly_z]),
            blended_rate,
            mo.md(
                "Required columns: `meter_name`, `start_date`, `end_date`, `kwh`. "
                "A `cost` column may also be included. Exact billing dates are preferred."
            ),
        ]
    )
    controls
    return (
        anomaly_z,
        blended_rate,
        cdd_base,
        controls,
        cooling_threshold,
        humidity_threshold,
        latitude,
        longitude,
        utility_upload,
    )


@app.cell
def _(io, mo, pd, utility_upload):
    mo.stop(not utility_upload.value, mo.md("### Upload the utility CSV/XLSX above to begin."))

    uploaded_file = utility_upload.value[0]
    uploaded_name = uploaded_file.name
    uploaded_bytes = uploaded_file.contents

    if uploaded_name.lower().endswith(".csv"):
        utility_raw = pd.read_csv(io.BytesIO(uploaded_bytes))
    elif uploaded_name.lower().endswith((".xlsx", ".xls")):
        utility_raw = pd.read_excel(io.BytesIO(uploaded_bytes))
    else:
        raise ValueError("Upload a CSV or Excel file.")

    required_columns = {"meter_name", "start_date", "end_date", "kwh"}
    missing_columns = required_columns - set(utility_raw.columns)
    if missing_columns:
        raise ValueError(f"Missing required columns: {sorted(missing_columns)}")

    mo.md(f"**Loaded:** `{uploaded_name}` — {len(utility_raw):,} rows")
    return uploaded_bytes, uploaded_file, uploaded_name, utility_raw


@app.cell
def _(pd, utility_raw):
    plant_lookup = {
        "North Part 1": "North",
        "North Part 2": "North",
        "North Part 3": "North",
        "South TES 1": "South",
        "South TES 2": "South",
    }

    def normalize_meter_name(x):
        return (
            str(x)
            .upper()
            .replace("\xa0", " ")
            .replace("*", "")
            .replace("  ", " ")
            .strip()
        )

    def map_meter_name(name):
        n = normalize_meter_name(name)
        if "NORTH TES" in n and "PART 1" in n:
            return "North Part 1"
        if "NORTH TES" in n and "PART 2" in n:
            return "North Part 2"
        if "NORTH TES" in n and "PART 3" in n:
            return "North Part 3"
        if "SOUTH TES" in n and ("- 2" in n or "TES 2" in n or "CHILL PLT 2" in n):
            return "South TES 2"
        if "SOUTH TES" in n:
            return "South TES 1"
        return "Unmapped"

    utility = utility_raw.copy()
    utility["meter_group"] = utility["meter_name"].apply(map_meter_name)
    utility["plant"] = utility["meter_group"].map(plant_lookup)
    utility["start_date"] = pd.to_datetime(utility["start_date"])
    utility["end_date"] = pd.to_datetime(utility["end_date"])
    utility["kwh"] = pd.to_numeric(utility["kwh"], errors="coerce")
    if "cost" in utility.columns:
        utility["cost"] = pd.to_numeric(utility["cost"], errors="coerce")
    utility["days"] = (utility["end_date"] - utility["start_date"]).dt.days + 1
    utility["kwh_per_day"] = utility["kwh"] / utility["days"]

    unmapped_names = utility.loc[utility["meter_group"] == "Unmapped", "meter_name"].drop_duplicates()
    return map_meter_name, normalize_meter_name, plant_lookup, unmapped_names, utility


@app.cell
def _(mo, unmapped_names, utility):
    mapping_table = utility[["meter_name", "meter_group", "plant"]].drop_duplicates()
    warning = (
        mo.callout(
            mo.vstack(
                [
                    mo.md("**Unmapped meter names detected:**"),
                    mo.ui.table(unmapped_names.to_frame(name="meter_name"), selection=None),
                ]
            ),
            kind="warn",
        )
        if len(unmapped_names)
        else mo.callout("All meter names mapped successfully.", kind="success")
    )
    mo.vstack([mo.md("## 2. Meter mapping"), mo.ui.table(mapping_table, selection=None), warning])
    return mapping_table, warning


@app.cell
def _(pd, utility):
    meter_period = (
        utility[utility["plant"].notna()]
        .groupby(["plant", "meter_group", "start_date", "end_date"], as_index=False)
        .agg(kwh=("kwh", lambda s: s.sum(min_count=1)))
    )

    plant_period = (
        meter_period.groupby(["plant", "start_date", "end_date"], as_index=False)
        .agg(plant_kwh=("kwh", lambda s: s.sum(min_count=1)))
    )
    plant_period["days"] = (plant_period["end_date"] - plant_period["start_date"]).dt.days + 1
    plant_period["plant_kwh_per_day"] = plant_period["plant_kwh"] / plant_period["days"]

    meter_pivot = (
        meter_period.pivot_table(
            index=["plant", "start_date", "end_date"],
            columns="meter_group",
            values="kwh",
            aggfunc="sum",
        )
        .reset_index()
    )
    return meter_period, meter_pivot, plant_period


@app.cell
def _(meter_pivot, mo, plant_period):
    mo.vstack(
        [
            mo.md("## 3. Plant totals"),
            mo.md("**North = Part 1 + Part 2 + Part 3; South = TES 1 + TES 2.**"),
            mo.ui.table(plant_period, selection=None),
            mo.accordion({"Meter-period detail": mo.ui.table(meter_pivot, selection=None)}),
        ]
    )
    return


@app.cell
def _(cdd_base, cooling_threshold, humidity_threshold, latitude, longitude, mo, np, pd, plant_period, requests):
    timezone = "America/New_York"

    def fetch_openmeteo_hourly(latitude_value, longitude_value, start_date, end_date, timezone_value):
        url = "https://archive-api.open-meteo.com/v1/archive"
        params = {
            "latitude": latitude_value,
            "longitude": longitude_value,
            "start_date": pd.Timestamp(start_date).strftime("%Y-%m-%d"),
            "end_date": pd.Timestamp(end_date).strftime("%Y-%m-%d"),
            "hourly": "temperature_2m,relative_humidity_2m,dew_point_2m",
            "temperature_unit": "fahrenheit",
            "timezone": timezone_value,
        }
        response = requests.get(url, params=params, timeout=60)
        response.raise_for_status()
        data = response.json()
        if "hourly" not in data:
            raise RuntimeError(f"No hourly weather returned: {data}")
        frame = pd.DataFrame(data["hourly"])
        frame["time"] = pd.to_datetime(frame["time"])
        return frame.rename(
            columns={
                "temperature_2m": "temp_f",
                "relative_humidity_2m": "rh_pct",
                "dew_point_2m": "dewpoint_f",
            }
        )

    mo.stop(plant_period.empty, mo.md("No mapped plant periods available."))
    weather = fetch_openmeteo_hourly(
        latitude.value,
        longitude.value,
        plant_period["start_date"].min(),
        plant_period["end_date"].max(),
        timezone,
    )

    def f_to_c(f):
        return (f - 32.0) * 5.0 / 9.0

    def c_to_f(c):
        return c * 9.0 / 5.0 + 32.0

    def stull_wetbulb_c(temp_c, rh):
        rh_array = np.clip(np.asarray(rh, dtype=float), 1, 100)
        t_array = np.asarray(temp_c, dtype=float)
        return (
            t_array * np.arctan(0.151977 * np.sqrt(rh_array + 8.313659))
            + np.arctan(t_array + rh_array)
            - np.arctan(rh_array - 1.676331)
            + 0.00391838 * rh_array**1.5 * np.arctan(0.023101 * rh_array)
            - 4.686035
        )

    weather["wetbulb_f"] = c_to_f(stull_wetbulb_c(f_to_c(weather["temp_f"]), weather["rh_pct"]))
    weather["cooling_degree_hour_65"] = np.maximum(weather["temp_f"] - cdd_base.value, 0)
    weather["humid_cooling_hour"] = (
        (weather["temp_f"] >= cooling_threshold.value)
        & (weather["rh_pct"] >= humidity_threshold.value)
    ).astype(int)
    return c_to_f, f_to_c, fetch_openmeteo_hourly, stull_wetbulb_c, timezone, weather


@app.cell
def _(mo, weather):
    mo.callout(f"Downloaded {len(weather):,} hourly Open-Meteo weather records.", kind="success")
    return


@app.cell
def _(pd, plant_period, weather):
    def summarize_weather(start_date, end_date):
        start = pd.Timestamp(start_date)
        end_exclusive = pd.Timestamp(end_date) + pd.Timedelta(days=1)
        period_weather = weather[(weather["time"] >= start) & (weather["time"] < end_exclusive)].copy()
        if period_weather.empty:
            return pd.Series(dtype=float)
        return pd.Series(
            {
                "weather_hours": len(period_weather),
                "mean_temp_f": period_weather["temp_f"].mean(),
                "max_temp_f": period_weather["temp_f"].max(),
                "min_temp_f": period_weather["temp_f"].min(),
                "mean_rh_pct": period_weather["rh_pct"].mean(),
                "mean_dewpoint_f": period_weather["dewpoint_f"].mean(),
                "mean_wetbulb_f": period_weather["wetbulb_f"].mean(),
                "max_wetbulb_f": period_weather["wetbulb_f"].max(),
                "cdd65": period_weather["cooling_degree_hour_65"].sum() / 24.0,
                "humid_cooling_hours": period_weather["humid_cooling_hour"].sum(),
                "hours_above_80f": (period_weather["temp_f"] >= 80).sum(),
                "hours_above_90f": (period_weather["temp_f"] >= 90).sum(),
            }
        )

    weather_features = plant_period.apply(
        lambda r: summarize_weather(r["start_date"], r["end_date"]), axis=1
    )
    analysis = pd.concat(
        [plant_period.reset_index(drop=True), weather_features.reset_index(drop=True)], axis=1
    )
    analysis = analysis[analysis["plant_kwh"].notna()].copy()
    analysis["kwh_per_cdd"] = np.where(
        analysis["cdd65"] > 0, analysis["plant_kwh"] / analysis["cdd65"], np.nan
    )
    analysis["cdd65_per_day"] = analysis["cdd65"] / analysis["days"]
    return analysis, summarize_weather, weather_features


@app.cell
def _(LinearRegression, analysis, anomaly_z, mean_absolute_percentage_error, np, pd, r2_score):
    def fit_plant_model(d):
        modeled = d.copy().dropna(subset=["plant_kwh", "cdd65", "mean_wetbulb_f", "days"])
        if len(modeled) < 4:
            for col in ["expected_kwh", "residual_kwh", "residual_pct", "performance_index", "residual_z", "model_r2", "model_mape_pct"]:
                modeled[col] = np.nan
            modeled["model_features"] = "insufficient data"
            return modeled, None

        features = ["cdd65", "mean_wetbulb_f", "days"] if len(modeled) >= 8 else ["cdd65", "days"]
        X = modeled[features]
        y = modeled["plant_kwh"]
        model = LinearRegression().fit(X, y)
        pred = model.predict(X)
        modeled["expected_kwh"] = pred
        modeled["residual_kwh"] = modeled["plant_kwh"] - modeled["expected_kwh"]
        modeled["residual_pct"] = 100 * modeled["residual_kwh"] / modeled["expected_kwh"]
        modeled["performance_index"] = modeled["plant_kwh"] / modeled["expected_kwh"]
        residual_sd = modeled["residual_kwh"].std(ddof=1)
        modeled["residual_z"] = (
            (modeled["residual_kwh"] - modeled["residual_kwh"].mean()) / residual_sd
            if pd.notna(residual_sd) and residual_sd > 0
            else 0.0
        )
        modeled["model_r2"] = r2_score(y, pred)
        modeled["model_mape_pct"] = mean_absolute_percentage_error(y, pred) * 100
        modeled["model_features"] = ", ".join(features)
        return modeled, model

    modeled_parts = []
    plant_models = {}
    for _plant_name, _plant_data in analysis.groupby("plant"):
        _plant_out, _plant_model = fit_plant_model(_plant_data)
        modeled_parts.append(_plant_out)
        plant_models[_plant_name] = _plant_model

    results = pd.concat(modeled_parts, ignore_index=True).sort_values(["plant", "start_date"])
    results["flag"] = np.select(
        [results["residual_z"] >= anomaly_z.value, results["residual_z"] <= -anomaly_z.value],
        ["HIGH USE", "LOW USE"],
        default="Normal",
    )

    diagnostics = (
        results.groupby("plant")
        .agg(
            billing_periods=("plant_kwh", "count"),
            total_kwh=("plant_kwh", "sum"),
            mean_cdd65=("cdd65", "mean"),
            r2=("model_r2", "first"),
            mape_pct=("model_mape_pct", "first"),
            features=("model_features", "first"),
        )
        .reset_index()
    )

    flagged = results.loc[results["flag"] != "Normal", [
        "plant", "start_date", "end_date", "plant_kwh", "expected_kwh", "residual_kwh",
        "residual_pct", "performance_index", "cdd65", "mean_wetbulb_f", "residual_z", "flag"
    ]].copy()
    return diagnostics, fit_plant_model, flagged, modeled_parts, plant_models, results


@app.cell
def _(analysis, mo, plt, results):
    energy_figs = []
    for _plant_name, _plant_data in results.groupby("plant"):
        _d = _plant_data.sort_values("start_date")
        _fig, _ax = plt.subplots(figsize=(9, 4.5))
        _ax.plot(_d["start_date"], _d["plant_kwh"], marker="o", label="Actual kWh")
        _ax.plot(_d["start_date"], _d["expected_kwh"], marker="o", label="Weather-normalized expected kWh")
        _ax.set_title(f"{_plant_name} Plant - Actual vs Expected")
        _ax.set_xlabel("Billing period")
        _ax.set_ylabel("kWh")
        _ax.legend()
        _ax.grid(alpha=0.25)
        _fig.tight_layout()
        energy_figs.append(mo.vstack([mo.md(f"### {_plant_name} Plant"), _fig]))

    cdd_fig, cdd_ax = plt.subplots(figsize=(9, 4.5))
    for _plant_name, _plant_data in analysis.groupby("plant"):
        _d = _plant_data.sort_values("start_date")
        cdd_ax.plot(_d["start_date"], _d["kwh_per_cdd"], marker="o", label=_plant_name)
    cdd_ax.set_title("North vs South - Plant kWh per CDD65")
    cdd_ax.set_xlabel("Billing period")
    cdd_ax.set_ylabel("kWh / CDD65")
    cdd_ax.legend()
    cdd_ax.grid(alpha=0.25)
    cdd_fig.tight_layout()

    mo.vstack([mo.md("## 4. Weather-normalized performance"), cdd_fig] + energy_figs)
    return cdd_ax, cdd_fig, energy_figs


@app.cell
def _(diagnostics, mo, results):
    valid_results = results[
        results["plant_kwh"].notna()
        & results["expected_kwh"].notna()
        & results["performance_index"].notna()
    ].copy()
    latest = valid_results.sort_values("start_date").groupby("plant").tail(1).copy()

    latest_display = latest[[
        "plant", "end_date", "plant_kwh", "expected_kwh", "performance_index", "residual_pct", "flag"
    ]].rename(columns={
        "plant": "Plant", "end_date": "Period End", "plant_kwh": "Actual kWh",
        "expected_kwh": "Expected kWh", "performance_index": "Performance Index",
        "residual_pct": "Residual %", "flag": "Flag"
    })

    mo.vstack([
        mo.md("## 5. Latest KPIs and model diagnostics"),
        mo.ui.table(latest_display, selection=None),
        mo.md("**R-squared:** fraction of historical variation explained by the model. **MAPE:** average absolute percentage prediction error; it can become misleading when actual kWh is near zero."),
        mo.ui.table(diagnostics, selection=None),
    ])
    return latest, latest_display, valid_results


@app.cell
def _(PCA, StandardScaler, analysis, np, pd):
    pca_weather_features = [
        "mean_temp_f", "max_temp_f", "mean_rh_pct", "mean_dewpoint_f", "mean_wetbulb_f",
        "max_wetbulb_f", "cdd65_per_day",
    ]
    pca_base = analysis.copy()
    pca_base["humid_cooling_hours_per_day"] = pca_base["humid_cooling_hours"] / pca_base["days"]
    pca_base["hours_above_80f_per_day"] = pca_base["hours_above_80f"] / pca_base["days"]
    pca_base["hours_above_90f_per_day"] = pca_base["hours_above_90f"] / pca_base["days"]
    pca_weather_features += [
        "humid_cooling_hours_per_day", "hours_above_80f_per_day", "hours_above_90f_per_day"
    ]

    weather_periods = (
        pca_base[["start_date", "end_date"] + pca_weather_features]
        .drop_duplicates(subset=["start_date", "end_date"])
        .sort_values("start_date")
        .dropna(subset=pca_weather_features)
        .reset_index(drop=True)
    )

    if len(weather_periods) >= 4:
        pca_scaler = StandardScaler()
        scaled_weather = pca_scaler.fit_transform(weather_periods[pca_weather_features])
        n_components = min(3, len(pca_weather_features), len(weather_periods))
        pca_model = PCA(n_components=n_components)
        score_array = pca_model.fit_transform(scaled_weather)
        component_names = [f"PC{i+1}" for i in range(n_components)]
        pca_scores = weather_periods[["start_date", "end_date"]].copy()
        for index, component_name in enumerate(component_names):
            pca_scores[component_name] = score_array[:, index]
        pca_summary = pd.DataFrame({
            "Component": component_names,
            "Explained_Variance_pct": pca_model.explained_variance_ratio_ * 100,
            "Cumulative_Variance_pct": np.cumsum(pca_model.explained_variance_ratio_) * 100,
        })
        pca_loadings = pd.DataFrame(
            pca_model.components_.T,
            index=pca_weather_features,
            columns=component_names,
        ).reset_index().rename(columns={"index": "Weather_Variable"})
    else:
        pca_scaler = None
        pca_model = None
        pca_scores = pd.DataFrame()
        pca_summary = pd.DataFrame()
        pca_loadings = pd.DataFrame()
    return (
        component_names if len(weather_periods) >= 4 else [],
        n_components if len(weather_periods) >= 4 else 0,
        pca_base,
        pca_loadings,
        pca_model,
        pca_scaler,
        pca_scores,
        pca_summary,
        pca_weather_features,
        score_array if len(weather_periods) >= 4 else np.empty((0, 0)),
        scaled_weather if len(weather_periods) >= 4 else np.empty((0, 0)),
        weather_periods,
    )


@app.cell
def _(PCA, StandardScaler, np, pd, pca_base, pca_scores, pca_summary):
    # Sensitivity test: repeat PCA with a leaner feature set so temperature
    # severity does not receive multiple redundant representations.
    reduced_pca_features = [
        "mean_temp_f",
        "max_temp_f",
        "mean_rh_pct",
        "mean_wetbulb_f",
        "max_wetbulb_f",
        "cdd65_per_day",
        "humid_cooling_hours_per_day",
    ]

    reduced_weather_periods = (
        pca_base[["start_date", "end_date"] + reduced_pca_features]
        .drop_duplicates(subset=["start_date", "end_date"])
        .sort_values("start_date")
        .dropna(subset=reduced_pca_features)
        .reset_index(drop=True)
    )

    if len(reduced_weather_periods) >= 4:
        reduced_pca_scaler = StandardScaler()
        _reduced_scaled = reduced_pca_scaler.fit_transform(
            reduced_weather_periods[reduced_pca_features]
        )
        _reduced_n_components = min(3, len(reduced_pca_features), len(reduced_weather_periods))
        reduced_pca_model = PCA(n_components=_reduced_n_components)
        _reduced_score_array = reduced_pca_model.fit_transform(_reduced_scaled)
        _reduced_component_names = [f"PC{i+1}" for i in range(_reduced_n_components)]

        reduced_pca_scores = reduced_weather_periods[["start_date", "end_date"]].copy()
        for _idx, _component in enumerate(_reduced_component_names):
            reduced_pca_scores[_component] = _reduced_score_array[:, _idx]

        reduced_pca_summary = pd.DataFrame({
            "Component": _reduced_component_names,
            "Explained_Variance_pct": reduced_pca_model.explained_variance_ratio_ * 100,
            "Cumulative_Variance_pct": np.cumsum(reduced_pca_model.explained_variance_ratio_) * 100,
        })
        reduced_pca_loadings = pd.DataFrame(
            reduced_pca_model.components_.T,
            index=reduced_pca_features,
            columns=_reduced_component_names,
        ).reset_index().rename(columns={"index": "Weather_Variable"})

        # Compare the geometry of the first two PCA dimensions. This avoids
        # relying on component signs, which may flip without changing meaning.
        _paired = pca_scores[["start_date", "end_date", "PC1", "PC2"]].merge(
            reduced_pca_scores[["start_date", "end_date", "PC1", "PC2"]],
            on=["start_date", "end_date"],
            suffixes=("_full", "_reduced"),
            how="inner",
        ).sort_values("start_date").reset_index(drop=True)

        if len(_paired) >= 4:
            _full_xy = _paired[["PC1_full", "PC2_full"]].to_numpy(float)
            _reduced_xy = _paired[["PC1_reduced", "PC2_reduced"]].to_numpy(float)
            _full_dist = np.sqrt(((_full_xy[:, None, :] - _full_xy[None, :, :]) ** 2).sum(axis=2))
            _reduced_dist = np.sqrt(((_reduced_xy[:, None, :] - _reduced_xy[None, :, :]) ** 2).sum(axis=2))
            _tri = np.triu_indices(len(_paired), k=1)
            distance_geometry_correlation = float(
                np.corrcoef(_full_dist[_tri], _reduced_dist[_tri])[0, 1]
            )

            _neighbor_rows = []
            _k = min(3, len(_paired) - 1)
            for _i in range(len(_paired)):
                _full_order = np.argsort(_full_dist[_i])
                _reduced_order = np.argsort(_reduced_dist[_i])
                _full_neighbors = [int(x) for x in _full_order if x != _i][:_k]
                _reduced_neighbors = [int(x) for x in _reduced_order if x != _i][:_k]
                _overlap = len(set(_full_neighbors) & set(_reduced_neighbors)) / _k if _k else np.nan
                _neighbor_rows.append({
                    "Period_End": _paired.loc[_i, "end_date"],
                    "Nearest_3_Weather_Neighbor_Overlap_pct": 100 * _overlap,
                })
            pca_neighbor_stability = pd.DataFrame(_neighbor_rows)
            mean_neighbor_overlap_pct = float(
                pca_neighbor_stability["Nearest_3_Weather_Neighbor_Overlap_pct"].mean()
            )
        else:
            distance_geometry_correlation = np.nan
            mean_neighbor_overlap_pct = np.nan
            pca_neighbor_stability = pd.DataFrame()

        _full_first_two = float(pca_summary["Explained_Variance_pct"].head(2).sum()) if not pca_summary.empty else np.nan
        _reduced_first_two = float(reduced_pca_summary["Explained_Variance_pct"].head(2).sum())
        pca_sensitivity_summary = pd.DataFrame([
            {
                "PCA_Version": "Full weather feature set",
                "Feature_Count": 10,
                "PC1_Explained_pct": float(pca_summary.loc[0, "Explained_Variance_pct"]) if not pca_summary.empty else np.nan,
                "PC1_PC2_Cumulative_pct": _full_first_two,
            },
            {
                "PCA_Version": "Reduced weather feature set",
                "Feature_Count": len(reduced_pca_features),
                "PC1_Explained_pct": float(reduced_pca_summary.loc[0, "Explained_Variance_pct"]),
                "PC1_PC2_Cumulative_pct": _reduced_first_two,
            },
        ])
        pca_sensitivity_metrics = pd.DataFrame([
            {
                "Metric": "Pairwise weather-distance correlation",
                "Value": distance_geometry_correlation,
                "Interpretation": "Closer to 1.0 means the reduced PCA preserves the full PCA weather geometry.",
            },
            {
                "Metric": "Mean nearest-3 weather-neighbor overlap",
                "Value": mean_neighbor_overlap_pct / 100 if pd.notna(mean_neighbor_overlap_pct) else np.nan,
                "Interpretation": "Closer to 1.0 means the same months remain nearest weather analogs.",
            },
        ])
    else:
        reduced_pca_scaler = None
        reduced_pca_model = None
        reduced_pca_scores = pd.DataFrame()
        reduced_pca_summary = pd.DataFrame()
        reduced_pca_loadings = pd.DataFrame()
        pca_neighbor_stability = pd.DataFrame()
        pca_sensitivity_summary = pd.DataFrame()
        pca_sensitivity_metrics = pd.DataFrame()
        distance_geometry_correlation = np.nan
        mean_neighbor_overlap_pct = np.nan

    return (
        distance_geometry_correlation,
        mean_neighbor_overlap_pct,
        pca_neighbor_stability,
        pca_sensitivity_metrics,
        pca_sensitivity_summary,
        reduced_pca_features,
        reduced_pca_loadings,
        reduced_pca_model,
        reduced_pca_scaler,
        reduced_pca_scores,
        reduced_pca_summary,
        reduced_weather_periods,
    )


@app.cell
def _(mo, pca_loadings, pca_scores, pca_summary, pd, plt, results):
    if pca_scores.empty or "PC2" not in pca_scores.columns:
        pca_results = results.copy()
        pca_panel = mo.callout("Insufficient complete weather periods for PCA.", kind="warn")
    else:
        pca_results = results.merge(pca_scores, on=["start_date", "end_date"], how="left")
        pca_figs = []
        for _plant_name, _plant_data in pca_results.dropna(subset=["PC1", "PC2", "residual_pct"]).groupby("plant"):
            _d = _plant_data.sort_values("start_date").copy()
            _fig, _ax = plt.subplots(figsize=(8.5, 5.0))
            _sizes = 35 + 3 * _d["residual_pct"].abs().clip(upper=60)
            _sc = _ax.scatter(_d["PC1"], _d["PC2"], s=_sizes, c=_d["residual_pct"], alpha=0.75)
            _ax.axhline(0, linewidth=0.8, alpha=0.4)
            _ax.axvline(0, linewidth=0.8, alpha=0.4)
            _ax.set_title(f"{_plant_name} Plant - PCA Weather Regimes vs Energy Residual")
            _ax.set_xlabel(f"PC1 ({pca_summary.loc[0, 'Explained_Variance_pct']:.1f}% variance)")
            _ax.set_ylabel(f"PC2 ({pca_summary.loc[1, 'Explained_Variance_pct']:.1f}% variance)")
            _colorbar = _fig.colorbar(_sc, ax=_ax)
            _colorbar.set_label("Energy residual (%)")
            _ax.grid(alpha=0.2)
            _labels = pd.concat([
                _d.nlargest(min(5, len(_d)), "residual_pct", keep="all"),
                _d.nsmallest(min(3, len(_d)), "residual_pct", keep="all"),
            ]).drop_duplicates(subset=["start_date"])
            for _, _row in _labels.iterrows():
                _ax.annotate(
                    pd.to_datetime(_row["end_date"]).strftime("%b %Y"),
                    (_row["PC1"], _row["PC2"]),
                    xytext=(4, 4), textcoords="offset points", fontsize=8,
                )
            _fig.tight_layout()
            pca_figs.append(mo.vstack([mo.md(f"### {_plant_name}"), _fig]))

        first_two_variance = pca_summary["Explained_Variance_pct"].head(2).sum()
        pca_panel = mo.vstack([
            mo.md("## 6. PCA weather-regime analysis"),
            mo.md(
                f"The first two PCA components explain **{first_two_variance:.1f}%** of the variation in the selected weather variables. "
                "Points close together experienced similar combinations of temperature, humidity, wet-bulb, CDD, and hot/humid hours. "
                "Large differences in energy residual among nearby points are stronger candidates for operational review."
            ),
            mo.ui.table(pca_summary, selection=None),
            mo.accordion({"PCA loadings": mo.ui.table(pca_loadings, selection=None)}),
            *pca_figs,
        ])
    pca_panel
    return pca_panel, pca_results


@app.cell
def _(
    distance_geometry_correlation,
    mean_neighbor_overlap_pct,
    mo,
    pca_sensitivity_metrics,
    pca_sensitivity_summary,
    pd,
    plt,
    reduced_pca_loadings,
    reduced_pca_scores,
    reduced_pca_summary,
    results,
):
    if reduced_pca_scores.empty or "PC2" not in reduced_pca_scores.columns:
        pca_sensitivity_panel = mo.callout(
            "Insufficient complete weather periods for the reduced-feature PCA sensitivity test.",
            kind="warn",
        )
    else:
        _reduced_results = results.merge(
            reduced_pca_scores, on=["start_date", "end_date"], how="left"
        )
        _reduced_figs = []
        for _plant_name, _plant_data in _reduced_results.dropna(
            subset=["PC1", "PC2", "residual_pct"]
        ).groupby("plant"):
            _d = _plant_data.sort_values("start_date").copy()
            _fig, _ax = plt.subplots(figsize=(8.5, 5.0))
            _sizes = 35 + 3 * _d["residual_pct"].abs().clip(upper=60)
            _sc = _ax.scatter(
                _d["PC1"], _d["PC2"], s=_sizes, c=_d["residual_pct"], alpha=0.75
            )
            _ax.axhline(0, linewidth=0.8, alpha=0.4)
            _ax.axvline(0, linewidth=0.8, alpha=0.4)
            _ax.set_title(f"{_plant_name} Plant - Reduced PCA Weather Regimes")
            _ax.set_xlabel(
                f"PC1 ({reduced_pca_summary.loc[0, 'Explained_Variance_pct']:.1f}% variance)"
            )
            _ax.set_ylabel(
                f"PC2 ({reduced_pca_summary.loc[1, 'Explained_Variance_pct']:.1f}% variance)"
            )
            _cb = _fig.colorbar(_sc, ax=_ax)
            _cb.set_label("Energy residual (%)")
            _ax.grid(alpha=0.2)
            _labels = pd.concat([
                _d.nlargest(min(5, len(_d)), "residual_pct", keep="all"),
                _d.nsmallest(min(3, len(_d)), "residual_pct", keep="all"),
            ]).drop_duplicates(subset=["start_date"])
            for _, _row in _labels.iterrows():
                _ax.annotate(
                    pd.to_datetime(_row["end_date"]).strftime("%b %Y"),
                    (_row["PC1"], _row["PC2"]),
                    xytext=(4, 4),
                    textcoords="offset points",
                    fontsize=8,
                )
            _fig.tight_layout()
            _reduced_figs.append(mo.vstack([mo.md(f"### {_plant_name}"), _fig]))

        if pd.notna(distance_geometry_correlation) and distance_geometry_correlation >= 0.85:
            _geometry_text = "strong"
        elif pd.notna(distance_geometry_correlation) and distance_geometry_correlation >= 0.70:
            _geometry_text = "moderate"
        else:
            _geometry_text = "limited"

        if pd.notna(mean_neighbor_overlap_pct) and mean_neighbor_overlap_pct >= 70:
            _neighbor_text = "strong"
        elif pd.notna(mean_neighbor_overlap_pct) and mean_neighbor_overlap_pct >= 50:
            _neighbor_text = "moderate"
        else:
            _neighbor_text = "limited"

        pca_sensitivity_panel = mo.vstack([
            mo.md("## 7. PCA sensitivity test"),
            mo.md(
                "The full PCA is repeated with a reduced weather feature set: mean temperature, maximum temperature, "
                "mean relative humidity, mean and maximum wet-bulb, CDD65/day, and humid-cooling-hours/day. "
                "This tests whether the weather-regime conclusions depend too heavily on redundant heat variables."
            ),
            mo.ui.table(pca_sensitivity_summary, selection=None),
            mo.ui.table(pca_sensitivity_metrics, selection=None),
            mo.callout(
                f"Weather-regime geometry agreement is **{_geometry_text}** "
                f"(pairwise-distance correlation = {distance_geometry_correlation:.2f}). "
                f"Nearest-weather-month agreement is **{_neighbor_text}** "
                f"(mean nearest-3 overlap = {mean_neighbor_overlap_pct:.0f}%). "
                "If both are high, the PCA conclusions are robust to removal of redundant weather variables.",
                kind="info",
            ),
            mo.accordion({
                "Reduced PCA explained variance": mo.ui.table(reduced_pca_summary, selection=None),
                "Reduced PCA loadings": mo.ui.table(reduced_pca_loadings, selection=None),
            }),
            *_reduced_figs,
            mo.md(
                "**Important:** the HIGH USE / LOW USE flags do not change in this sensitivity test because they come from "
                "the weather-normalized energy model, not PCA. The sensitivity test asks whether the *weather analogs* "
                "remain similar when redundant PCA inputs are removed."
            ),
        ])
    pca_sensitivity_panel
    return pca_sensitivity_panel


@app.cell
def _(meter_period, np, pd):
    expected_meters = {"North": 3, "South": 2}
    meter_detail = meter_period.copy()
    period_completeness = (
        meter_detail[meter_detail["kwh"].notna()]
        .groupby(["plant", "start_date", "end_date"], as_index=False)
        .agg(
            meters_present=("meter_group", "nunique"),
            complete_plant_kwh=("kwh", lambda s: s.sum(min_count=1)),
        )
    )
    period_completeness["meters_expected"] = period_completeness["plant"].map(expected_meters)
    period_completeness["complete_period"] = (
        period_completeness["meters_present"] == period_completeness["meters_expected"]
    )
    meter_detail = meter_detail.merge(
        period_completeness, on=["plant", "start_date", "end_date"], how="left"
    )
    meter_detail["meter_share_pct"] = np.where(
        meter_detail["complete_period"] & meter_detail["kwh"].notna(),
        100 * meter_detail["kwh"] / meter_detail["complete_plant_kwh"],
        np.nan,
    )
    complete_meter_detail = meter_detail[
        meter_detail["complete_period"] & meter_detail["kwh"].notna()
    ].copy()
    meter_summary = (
        complete_meter_detail.groupby(["plant", "meter_group"], as_index=False)
        .agg(
            Complete_Period_kWh=("kwh", "sum"),
            Complete_Billing_Periods=("start_date", "nunique"),
        )
    )
    complete_plant_totals = meter_summary.groupby("plant")["Complete_Period_kWh"].transform("sum")
    meter_summary["Contribution_pct"] = 100 * meter_summary["Complete_Period_kWh"] / complete_plant_totals
    meter_summary = meter_summary[[
        "plant", "meter_group", "Complete_Billing_Periods", "Complete_Period_kWh", "Contribution_pct"
    ]]
    return complete_meter_detail, complete_plant_totals, expected_meters, meter_detail, meter_summary, period_completeness


@app.cell
def _(blended_rate, pd, results):
    cost_results = results.copy()
    cost_results["positive_residual_kwh"] = cost_results["residual_kwh"].clip(lower=0)
    cost_results["estimated_excess_energy_cost"] = (
        cost_results["positive_residual_kwh"] * blended_rate.value
    )
    cost_summary = (
        cost_results.groupby("plant")
        .agg(
            positive_residual_kwh=("positive_residual_kwh", "sum"),
            estimated_excess_energy_cost=("estimated_excess_energy_cost", "sum"),
        )
        .reset_index()
    )

    chillers = pd.DataFrame([
        ["South", "CH-1", 1000, 0.5562],
        ["South", "CH-2", 1000, 0.5562],
        ["South", "CH-3", 1000, 0.5562],
        ["North", "CH-1", 1000, 0.5700],
        ["North", "CH-2", 1000, 0.5240],
        ["North", "CH-3", 1000, 0.6620],
    ], columns=["plant", "chiller", "rated_tons", "rated_kw_per_ton"])
    chillers["rated_full_load_kw"] = chillers["rated_tons"] * chillers["rated_kw_per_ton"]
    return chillers, cost_results, cost_summary


@app.cell
def _(cost_summary, flagged, meter_summary, mo):
    mo.vstack([
        mo.md("## 8. Investigation screens"),
        mo.md("### Flagged billing periods"),
        mo.ui.table(flagged, selection=None),
        mo.md("### Complete-period meter contribution"),
        mo.ui.table(meter_summary, selection=None),
        mo.md("Contribution percentages use only periods where every expected meter for that plant has kWh; they therefore sum to 100% within each plant."),
        mo.md("### Above-baseline cost screening"),
        mo.ui.table(cost_summary, selection=None),
        mo.callout("Above-baseline cost is a screening value, not verified savings.", kind="warn"),
    ])
    return


@app.cell
def _(io, pd, utility, meter_period, meter_pivot, plant_period, weather, results, diagnostics, flagged, chillers, cost_summary, pca_scores, pca_loadings, pca_summary, meter_summary, reduced_pca_scores, reduced_pca_loadings, reduced_pca_summary, pca_sensitivity_summary, pca_sensitivity_metrics, pca_neighbor_stability):
    excel_buffer = io.BytesIO()
    with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
        utility.to_excel(writer, sheet_name="Utility Raw Input", index=False)
        meter_period.to_excel(writer, sheet_name="Meter Period Detail", index=False)
        meter_pivot.to_excel(writer, sheet_name="Meter Pivot", index=False)
        plant_period.to_excel(writer, sheet_name="Plant Totals", index=False)
        weather.to_excel(writer, sheet_name="Hourly Weather", index=False)
        results.to_excel(writer, sheet_name="Performance Results", index=False)
        diagnostics.to_excel(writer, sheet_name="Model Diagnostics", index=False)
        flagged.to_excel(writer, sheet_name="Flagged Periods", index=False)
        meter_summary.to_excel(writer, sheet_name="Meter Contribution", index=False)
        chillers.to_excel(writer, sheet_name="Chiller Benchmarks", index=False)
        cost_summary.to_excel(writer, sheet_name="Cost Summary", index=False)
        if not pca_scores.empty:
            pca_scores.to_excel(writer, sheet_name="PCA Scores Full", index=False)
            pca_loadings.to_excel(writer, sheet_name="PCA Loadings Full", index=False)
            pca_summary.to_excel(writer, sheet_name="PCA Summary Full", index=False)
        if not reduced_pca_scores.empty:
            reduced_pca_scores.to_excel(writer, sheet_name="PCA Scores Reduced", index=False)
            reduced_pca_loadings.to_excel(writer, sheet_name="PCA Load Reduced", index=False)
            reduced_pca_summary.to_excel(writer, sheet_name="PCA Summary Reduced", index=False)
        if not pca_sensitivity_summary.empty:
            pca_sensitivity_summary.to_excel(writer, sheet_name="PCA Sensitivity", index=False)
        if not pca_sensitivity_metrics.empty:
            pca_sensitivity_metrics.to_excel(writer, sheet_name="PCA Robustness", index=False)
        if not pca_neighbor_stability.empty:
            pca_neighbor_stability.to_excel(writer, sheet_name="PCA Neighbor Stability", index=False)
    excel_bytes = excel_buffer.getvalue()
    return excel_buffer, excel_bytes


@app.cell
def _(
    Document,
    Inches,
    OxmlElement,
    Pt,
    WD_ALIGN_PARAGRAPH,
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT,
    chillers,
    cost_summary,
    diagnostics,
    flagged,
    io,
    latest,
    meter_summary,
    np,
    pca_summary,
    pca_sensitivity_summary,
    pca_sensitivity_metrics,
    reduced_pca_summary,
    pd,
    qn,
    valid_results,
):
    def set_cell_text(cell, text, bold=False, size=9):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(size)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def add_dataframe_table(doc, df, max_rows=40):
        if df is None or len(df) == 0:
            doc.add_paragraph("No data available.")
            return
        display_df = df.head(max_rows).copy()
        table = doc.add_table(rows=1, cols=len(display_df.columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for index, column in enumerate(display_df.columns):
            set_cell_text(table.rows[0].cells[index], column, bold=True, size=8)
        for _, row in display_df.iterrows():
            cells = table.add_row().cells
            for index, value in enumerate(row):
                if pd.isna(value):
                    text = ""
                elif isinstance(value, (float, np.floating)):
                    text = f"{value:,.2f}"
                else:
                    text = str(value)
                set_cell_text(cells[index], text, size=8)
        doc.add_paragraph()

    report_doc = Document()
    report_doc.sections[0].top_margin = Inches(0.7)
    report_doc.sections[0].bottom_margin = Inches(0.7)
    report_doc.sections[0].left_margin = Inches(0.75)
    report_doc.sections[0].right_margin = Inches(0.75)
    report_doc.styles["Normal"].font.name = "Arial"
    report_doc.styles["Normal"].font.size = Pt(10.5)
    report_doc.styles["Heading 1"].font.name = "Arial"
    report_doc.styles["Heading 1"].font.size = Pt(14)
    report_doc.styles["Heading 2"].font.name = "Arial"
    report_doc.styles["Heading 2"].font.size = Pt(12)

    title = report_doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ERAU Central Chiller Plant Performance Report")
    title_run.bold = True
    title_run.font.size = Pt(18)

    report_end = pd.to_datetime(valid_results["end_date"]).max().date()
    subtitle = report_doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Reporting through {report_end}")

    report_doc.add_heading("Executive Summary", level=1)
    if len(latest) >= 2:
        comparison = latest.set_index("plant")
        if "North" in comparison.index and "South" in comparison.index:
            report_doc.add_paragraph(
                f"In the latest modeled period, North had a Performance Index of {comparison.loc['North', 'performance_index']:.2f} "
                f"and South had a Performance Index of {comparison.loc['South', 'performance_index']:.2f}. "
                "Values above 1.00 indicate use above the weather-normalized expectation; values below 1.00 indicate use below expectation."
            )
    for _plant_name, _group in valid_results.groupby("plant"):
        _high_count = int((_group["flag"] == "HIGH USE").sum())
        report_doc.add_paragraph(f"{_plant_name} had {_high_count} billing period(s) flagged as unusually high use.")
    for _, _row in cost_summary.iterrows():
        report_doc.add_paragraph(
            f"{_row['plant']} positive-residual energy represents approximately ${_row['estimated_excess_energy_cost']:,.0f} at the configured blended rate. "
            "This is an investigation screen, not verified savings."
        )
    if not pca_summary.empty:
        report_doc.add_paragraph(
            f"The first two PCA weather components explain {pca_summary['Explained_Variance_pct'].head(2).sum():.1f}% of variation in the selected weather variables."
        )

    report_doc.add_heading("Latest Plant KPIs", level=1)
    latest_table = latest[["plant", "end_date", "plant_kwh", "expected_kwh", "performance_index", "residual_pct", "flag"]].copy()
    add_dataframe_table(report_doc, latest_table)

    report_doc.add_heading("Model Diagnostics", level=1)
    add_dataframe_table(report_doc, diagnostics)
    report_doc.add_paragraph(
        "R-squared indicates how much historical variation is explained by the model. MAPE indicates average absolute percentage error, but can become unstable when actual kWh approaches zero."
    )

    report_doc.add_heading("PCA Weather-Regime Analysis", level=1)
    if not pca_summary.empty:
        add_dataframe_table(report_doc, pca_summary)
        report_doc.add_paragraph(
            "PCA is derived from weather variables only. Months close together in PCA space experienced similar overall weather. "
            "Large differences in energy residual among nearby months are candidates for operational review because weather is less likely to explain the difference."
        )

    report_doc.add_heading("PCA Sensitivity Test", level=1)
    if not pca_sensitivity_summary.empty:
        add_dataframe_table(report_doc, pca_sensitivity_summary)
        if not pca_sensitivity_metrics.empty:
            add_dataframe_table(report_doc, pca_sensitivity_metrics)
        report_doc.add_paragraph(
            "The PCA was repeated with a reduced feature set to test whether redundant temperature-related variables materially change the weather-regime structure. "
            "Strong agreement between the full and reduced PCA indicates that the interpretation is robust rather than being driven by one particular feature set. "
            "The energy anomaly flags are unchanged because they are generated by the weather-normalized energy model, not by PCA."
        )

    report_doc.add_heading("Flagged Billing Periods", level=1)
    add_dataframe_table(report_doc, flagged)
    report_doc.add_heading("Meter Contribution Summary", level=1)
    add_dataframe_table(report_doc, meter_summary)
    report_doc.add_heading("Above-Baseline Energy Cost Screening", level=1)
    add_dataframe_table(report_doc, cost_summary)
    report_doc.add_heading("Rated Chiller Reference", level=1)
    add_dataframe_table(report_doc, chillers)

    report_doc.add_heading("Recommended Next Steps", level=1)
    recommendations = [
        "Validate unusual South Plant historical meter periods before treating them as operational performance.",
        "Review high-use periods against plant schedules, maintenance, staging, controls, and TES operation.",
        "Track Performance Index monthly to identify sustained changes.",
        "Use exact utility billing dates when available.",
        "Add BAS plant kW, chilled-water flow, CHWS/CHWR temperatures, chiller status/load, pump/tower operation, and TES status to enable actual kW/ton and operating-regime analysis.",
    ]
    for item in recommendations:
        report_doc.add_paragraph(item, style="List Number")

    report_doc.add_heading("Methodology Notes", level=1)
    for item in [
        "North Plant kWh = North Part 1 + Part 2 + Part 3.",
        "South Plant kWh = South TES 1 + South TES 2.",
        "Weather features are calculated from hourly Open-Meteo historical data.",
        "Performance Index = Actual kWh / Weather-Normalized Expected kWh.",
        "Cost-only months with no kWh are excluded from performance modeling.",
        "PCA uses standardized weather variables only and is exploratory.",
        "A reduced-feature PCA sensitivity test checks whether weather-regime conclusions persist after removing redundant weather variables.",
        "Results are screening indicators and do not represent measured chiller kW/ton.",
    ]:
        report_doc.add_paragraph(item, style="List Bullet")

    word_buffer = io.BytesIO()
    report_doc.save(word_buffer)
    word_bytes = word_buffer.getvalue()
    return add_dataframe_table, recommendations, report_doc, report_end, set_cell_text, word_buffer, word_bytes


@app.cell
def _(excel_bytes, mo, word_bytes):
    mo.vstack([
        mo.md("## 9. Download results"),
        mo.hstack([
            mo.download(
                data=excel_bytes,
                filename="ERAU_chiller_utility_weather_analysis.xlsx",
                label="Download Excel analysis",
            ),
            mo.download(
                data=word_bytes,
                filename="ERAU_chiller_performance_report.docx",
                label="Download Word report",
            ),
        ]),
        mo.md(
            "### What this analysis can and cannot tell you\n"
            "It can screen weather-normalized plant energy, anomalies, meter contributions, PCA weather regimes, and above-baseline cost. "
            "It cannot directly determine actual cooling tons, actual chiller/plant kW/ton, individual chiller sequencing efficiency, or pump/tower auxiliary efficiency without BAS data."
        ),
    ])
    return


if __name__ == "__main__":
    app.run()
